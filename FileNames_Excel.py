import os
from datetime import datetime
import pandas as pd
import streamlit as st
from pathlib import Path
from openpyxl.styles import Font, Alignment
import pyperclip

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from io import BytesIO


# Helper function to get only files (exclude folders)
def get_files(path):
    items = []
    for root, dirs, files in os.walk(path):
        for name in files:  # Only include files, exclude folders
            full_path = os.path.join(root, name)
            modified_time = datetime.fromtimestamp(os.path.getmtime(full_path)).strftime('%Y-%m-%d')
            items.append((name, modified_time, full_path))  # Include full path for preview
    return items

# Function to generate Word file
def generate_word(categories, output_path):
    doc = Document()

    # Create a table with two columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # Ensure 'Table Grid' style for visible borders

    # Set header cells
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category / File Name'
    hdr_cells[1].text = 'Last Modified'

    # Apply bold formatting to header cells
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True  # Make text bold
                run.font.size = Pt(12)  # Optional: set font size for clarity

    # Iterate over each category and file item
    for category, items in categories.items():
        if items:
            # Add a row for the category
            category_row = table.add_row().cells
            category_row[0].text = category
            category_row[1].text = ""  # Leave second cell blank
            category_row[0].paragraphs[0].runs[0].font.bold = True  # Bold category name

            # Add a row for each file in the category
            for name, modified_time in items:
                file_row = table.add_row().cells
                file_row[0].text = name
                file_row[1].text = modified_time

                # Optional: Set cell styles explicitly
                for cell in file_row:
                    cell.vertical_alignment = True  # Align text vertically at top

    # Save the document to the output path
    doc.save(output_path)
    print(f"Word document created at {output_path}")
    return output_path

# Function to generate PDF file
def generate_pdf(categories, output_path):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4)
    
    # Define the style for bold text and normal text
    styles = getSampleStyleSheet()
    bold_style = styles["Heading4"]
    bold_style.fontSize = 10
    bold_style.leading = 12
    bold_style.spaceAfter = 6
    bold_style.textColor = colors.black
    normal_style = styles["BodyText"]
    normal_style.fontSize = 10
    normal_style.leading = 12

    # Define the table data with a header row
    data = [
        [Paragraph("Category / File Name", bold_style), Paragraph("Last Modified", bold_style)]
    ]  # Header row with bold style

    # Add category and file rows
    for category, items in categories.items():
        if items:
            # Add a bold row for each category
            data.append([Paragraph(category, bold_style), ""])  # Category row

            # Add each file name and modified date as a new row
            for name, modified_time in items:
                wrapped_name = Paragraph(name, normal_style)
                data.append([wrapped_name, modified_time])  # Each file in a new row

    # Create and style the table
    table = Table(data, colWidths=[250, 100])  # Adjust column widths for clarity
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.white),         # Header row background color
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),    # Header text color
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),                  # Align text to left
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),           # Font for all cells
        ('FONTSIZE', (0, 0), (-1, -1), 10),                   # Font size for all rows
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),               # Padding for header
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),        # Black gridlines for all cells
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),                  # Vertical alignment at the top for all cells
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),       # White background for data rows
    ]))

    # Build the PDF with the table
    elements = [table]
    pdf.build(elements)
    buffer.seek(0)
    with open(output_path, "wb") as f:
        f.write(buffer.read())
    
    print(f"PDF document saved to {output_path}")
    return buffer

# Function to generate Excel file
def generate_excel(categories, output_path):

    data = []

    for category, items in categories.items():
        if items:
            data.append({'Category / File Name': category, 'Last Modified': ''})
            for name, modified_time in items:
                data.append({'Category / File Name': f"   {name}", 'Last Modified': modified_time})

    df = pd.DataFrame(data)

    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:

        df.to_excel(writer, index=False, sheet_name='Files', startrow=1)
        workbook = writer.book
        worksheet = writer.sheets['Files']
        headers = ['Category / File Name', 'Last Modified']

        for col_num, value in enumerate(headers, 1):
            cell = worksheet.cell(row=1, column=col_num, value=value)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='left')
            
        worksheet.column_dimensions['A'].width = 50
        worksheet.column_dimensions['B'].width = 20
        
        for index, row in df.iterrows():
            category_cell = worksheet.cell(row=index + 2, column=1, value=row['Category / File Name'])
            category_cell.alignment = Alignment(horizontal='left', vertical='top')
            last_modified_cell = worksheet.cell(row=index + 2, column=2, value=row['Last Modified'])
            last_modified_cell.alignment = Alignment(horizontal='left')

            if not row['Last Modified']:
                category_cell.font = Font(bold=True, size=12)

    return output_path

# Streamlit UI
st.title("File Categorization Tool")
directory_path = st.text_input("Enter the directory path:", "")

# Initialize session state for checkboxes (if not already present)
if 'generate_excel_option' not in st.session_state:
    st.session_state.generate_excel_option = False
if 'generate_word_option' not in st.session_state:
    st.session_state.generate_word_option = False
if 'generate_pdf_option' not in st.session_state:
    st.session_state.generate_pdf_option = False

# Initialize session state to hold generated file buffers (if not already present)
if 'generated_files' not in st.session_state:
    st.session_state.generated_files = {}

# Checkboxes to select file types for generation (persistent with session state)
st.session_state.generate_excel_option = st.checkbox("Generate Excel", value=st.session_state.generate_excel_option)
st.session_state.generate_word_option = st.checkbox("Generate Word", value=st.session_state.generate_word_option)
st.session_state.generate_pdf_option = st.checkbox("Generate PDF", value=st.session_state.generate_pdf_option)

if directory_path and Path(directory_path).exists():
    items = get_files(directory_path)
    
    if items:
        categories = ["CONTRACTUAL", "ARCHITECTURAL", "STRUCTURAL", "SERVICES", "SAFETY"]
        category_selection = {}

        st.write("### Assign Categories")
        for index, item in enumerate(items):
            name, modified_time, full_path = item
            cols = st.columns([3, 1])
            
            with cols[0]:
                if st.button(f"{index+1} {name}", key=f"copy_button_{index}"):
                    pyperclip.copy(full_path)
                    st.success(f"Path copied to clipboard: {full_path}")
            with cols[1]:
                category = st.selectbox("Select category", options=categories, key=f"selectbox_{index}", label_visibility="collapsed")
                category_selection[name] = (modified_time, category)

        # Generate files if not already generated
        if st.button("Generate Selected Files"):

            # Check if any file generation is needed
            categorized_data = {cat: [] for cat in categories}
            for name, (modified_time, category) in category_selection.items():
                categorized_data[category].append((name, modified_time))

            # Generate files only if they haven't been generated already
            if st.session_state.generate_excel_option and 'excel' not in st.session_state.generated_files:
                output_excel_buffer = generate_excel(categorized_data, "categorized_files.xlsx")
                st.session_state.generated_files['excel'] = output_excel_buffer

            if st.session_state.generate_word_option and 'word' not in st.session_state.generated_files:
                output_word_path = "categorized_files.docx"
                word_path = generate_word(categorized_data, output_word_path)
                st.session_state.generated_files['word'] = word_path

            if st.session_state.generate_pdf_option and 'pdf' not in st.session_state.generated_files:
                output_pdf_path = "categorized_files.pdf"
                pdf_buffer = generate_pdf(categorized_data, output_pdf_path)
                st.session_state.generated_files['pdf'] = pdf_buffer

        # Display download buttons for each generated file if available and selected
        if st.session_state.generate_excel_option and 'excel' in st.session_state.generated_files:
            st.download_button("Download Excel", data=st.session_state.generated_files['excel'], file_name="categorized_files.xlsx")

        if st.session_state.generate_word_option and 'word' in st.session_state.generated_files:
            st.download_button("Download Word", data=open(st.session_state.generated_files['word'], "rb"), file_name="categorized_files.docx")

        if st.session_state.generate_pdf_option and 'pdf' in st.session_state.generated_files:
            st.download_button("Download PDF", data=st.session_state.generated_files['pdf'], file_name="categorized_files.pdf")

    else:
        st.warning("No files found in the selected directory.")
else:
    st.info("Please enter a valid directory path to start categorizing files.")

